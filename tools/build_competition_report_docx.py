#!/usr/bin/env python3
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "competition_report_final.md"
OUT = ROOT / "docs" / "competition_report_final.docx"

FONT_BODY_CN = "SimSun"
FONT_BODY_CN_ALT = "宋体"
FONT_TITLE_CN = "SimSun"
FONT_TITLE_CN_ALT = "宋体"
FONT_EN = "SimSun"
BLACK = RGBColor(0, 0, 0)


def set_rfonts(run, east_asia=FONT_BODY_CN, ascii_font=FONT_EN):
    run.font.name = ascii_font
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)


def set_paragraph_font(paragraph, east_asia=FONT_BODY_CN, ascii_font=FONT_EN, size=None, bold=None):
    for run in paragraph.runs:
        set_rfonts(run, east_asia=east_asia, ascii_font=ascii_font)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        run.font.color.rgb = BLACK


def set_style_font(style, east_asia, ascii_font=FONT_EN, size=10.5, bold=False):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), ascii_font)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_width(table, width_dxa=9020):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    styles = doc.styles
    set_style_font(styles["Normal"], FONT_BODY_CN, size=10.5, bold=False)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    set_style_font(styles["Heading 1"], FONT_TITLE_CN, size=16, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], FONT_TITLE_CN, size=14, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], FONT_BODY_CN, size=12, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        if name in styles:
            set_style_font(styles[name], FONT_BODY_CN, size=10.5, bold=False)
            styles[name].paragraph_format.space_after = Pt(3)
            styles[name].paragraph_format.line_spacing = 1.15
    return doc


def add_text_runs(paragraph, text, east_asia=FONT_BODY_CN, size=10.5, bold=False):
    # Keep inline code readable but avoid Markdown backticks in final DOCX.
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    run = paragraph.add_run(text)
    set_rfonts(run, east_asia=east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = BLACK
    return run


def image_width(path):
    with Image.open(path) as img:
        w, h = img.size
    if w >= h:
        return Inches(5.9)
    return Inches(4.4)


def add_image(doc, md_dir, alt, rel):
    path = (md_dir / rel).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        add_text_runs(p, f"[图片缺失：{rel}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=image_width(path))
    set_rfonts(run)


def parse_table(lines, idx):
    rows = []
    while idx < len(lines):
        line = lines[idx]
        if not line.strip().startswith("|"):
            break
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade_cell(cell, "F2F2F2")
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(cell.text) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_rfonts(run, east_asia=FONT_BODY_CN)
                    run.font.size = Pt(9.5)
                    run.bold = r_idx == 0
                    run.font.color.rgb = BLACK
    doc.add_paragraph()


def mermaid_label(token):
    m = re.search(r"\[([^\]]+)\]", token)
    if m:
        return m.group(1)
    m = re.search(r"\{([^}]+)\}", token)
    if m:
        return m.group(1)
    token = re.sub(r"^[A-Za-z0-9_]+", "", token)
    return token.strip() or None


def add_mermaid_summary(doc, block):
    labels = []
    edges = []
    for line in block:
        line = line.strip()
        if not line or line.startswith("flowchart"):
            continue
        if "-->" in line:
            left, right = line.split("-->", 1)
            left_label = mermaid_label(left)
            right_label = mermaid_label(right)
            if left_label and right_label:
                edges.append((left_label, right_label))
                for item in (left_label, right_label):
                    if item not in labels:
                        labels.append(item)
    if labels:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_text_runs(p, "流程图：" + " → ".join(labels), east_asia=FONT_BODY_CN, size=10.5)
    elif edges:
        p = doc.add_paragraph()
        add_text_runs(p, "流程图：", east_asia=FONT_BODY_CN, size=10.5, bold=True)
    for left, right in edges[:12]:
        p = doc.add_paragraph(style="List Bullet")
        add_text_runs(p, f"{left} → {right}", east_asia=FONT_BODY_CN, size=10.5)


def build_doc():
    doc = setup_doc()
    md_dir = MD.parent
    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_block = []
    in_references = False

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip("`").strip()
                code_block = []
            else:
                if code_lang == "mermaid":
                    add_mermaid_summary(doc, code_block)
                else:
                    for code_line in code_block:
                        p = doc.add_paragraph()
                        add_text_runs(p, code_line, size=9.5)
                in_code = False
                code_lang = ""
            i += 1
            continue
        if in_code:
            code_block.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        img = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img:
            add_image(doc, md_dir, img.group(1), img.group(2))
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = add_text_runs(p, line[2:].strip(), east_asia=FONT_TITLE_CN, size=18, bold=True)
            i += 1
            continue

        if line.startswith("## "):
            in_references = "参考文献" in line
            p = doc.add_paragraph(style="Heading 1")
            add_text_runs(p, line[3:].strip(), east_asia=FONT_TITLE_CN, size=16, bold=True)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_text_runs(p, line[4:].strip(), east_asia=FONT_TITLE_CN, size=14, bold=True)
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_text_runs(p, line[5:].strip(), east_asia=FONT_BODY_CN, size=12, bold=True)
            i += 1
            continue

        if re.match(r"^- ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_text_runs(p, line[2:].strip(), east_asia=FONT_BODY_CN, size=10.5)
            i += 1
            continue

        m = re.match(r"^(\d+)\. (.*)", line)
        if m:
            if in_references:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                add_text_runs(p, line.strip(), east_asia=FONT_BODY_CN, size=9.5)
                i += 1
                continue
            p = doc.add_paragraph(style="List Number")
            add_text_runs(p, m.group(2).strip(), east_asia=FONT_BODY_CN, size=10.5)
            i += 1
            continue

        p = doc.add_paragraph()
        if re.match(r"^图\s*\d+", line) or line.startswith("图："):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_runs(p, line.strip(), east_asia=FONT_BODY_CN, size=10)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_text_runs(p, line.strip(), east_asia=FONT_BODY_CN, size=10.5)
        i += 1

    # Footer page number placeholder is intentionally omitted for simple submission style.
    doc.core_properties.title = MD.stem
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.subject = ""
    doc.core_properties.category = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    if len(sys.argv) == 3:
        MD = Path(sys.argv[1]).resolve()
        OUT = Path(sys.argv[2]).resolve()
    elif len(sys.argv) != 1:
        raise SystemExit("用法：build_competition_report_docx.py [输入.md 输出.docx]")
    build_doc()
